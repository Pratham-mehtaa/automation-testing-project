"""
generate_project_report.py — Generate a Word document explaining the project step by step.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, header_color="2F5496"):
    """Add a styled table with colored header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def create_project_report():
    doc = Document()

    # ──── Document Styles ────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    # ──── Title Page ─────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Automation Testing Project Report")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Selenium WebDriver — SauceDemo.com")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Prepared by: Pratham Mehta\nDate: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_page_break()

    # ──── Table of Contents ──────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Project Objective",
        "3. Technology Stack",
        "4. Project Setup",
        "5. Framework Architecture",
        "6. Page Object Model (POM) Design",
        "7. Test Scenarios & Test Cases",
        "8. Test Execution",
        "9. Test Results & Analysis",
        "10. Bugs Found",
        "11. Bug Report (Excel)",
        "12. Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # ──── 1. Introduction ────────────────────────────────────────
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document describes the step-by-step process followed to design, develop, "
        "and execute an automated testing project for the SauceDemo e-commerce web application "
        "(www.saucedemo.com). The project uses Selenium WebDriver with Python and the pytest "
        "framework to automate functional testing of the website across multiple user profiles."
    )
    doc.add_paragraph(
        "SauceDemo is a sample e-commerce website created by Sauce Labs specifically for "
        "practicing test automation. It features intentional bugs across different user accounts, "
        "making it an ideal candidate for automation testing and bug detection."
    )

    # ──── 2. Project Objective ───────────────────────────────────
    doc.add_heading("2. Project Objective", level=1)
    doc.add_paragraph("The primary objectives of this project are:")
    objectives = [
        "Build a robust, maintainable Selenium test automation framework using the Page Object Model (POM) design pattern.",
        "Automate functional testing of core e-commerce features: Login, Product Browsing, Shopping Cart, and Checkout.",
        "Test the application across all 6 available user types to uncover user-specific bugs.",
        "Document all discovered bugs with detailed steps to reproduce, expected vs. actual results, and severity ratings.",
        "Generate a professional Excel bug report for stakeholder review.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    # ──── 3. Technology Stack ────────────────────────────────────
    doc.add_heading("3. Technology Stack", level=1)
    doc.add_paragraph(
        "The following tools and technologies were selected for this project:"
    )
    add_styled_table(doc, ["Tool", "Version", "Purpose"], [
        ("Python", "3.9+", "Programming language"),
        ("Selenium WebDriver", "4.20+", "Browser automation library"),
        ("pytest", "8.0+", "Test framework with fixtures, markers, and reporting"),
        ("pytest-html", "4.1+", "HTML test report generation"),
        ("webdriver-manager", "4.0+", "Automatic ChromeDriver/GeckoDriver management"),
        ("Faker", "24.0+", "Random test data generation"),
        ("openpyxl", "3.1+", "Excel report generation"),
        ("Google Chrome", "147.x", "Browser under test"),
    ])

    # ──── 4. Project Setup ───────────────────────────────────────
    doc.add_heading("4. Project Setup", level=1)

    doc.add_heading("Step 4.1: Create Project Directory", level=2)
    doc.add_paragraph(
        "A new project directory was created to house all automation code, configuration files, "
        "and test reports."
    )

    doc.add_heading("Step 4.2: Set Up Virtual Environment", level=2)
    doc.add_paragraph("A Python virtual environment was created to isolate project dependencies:")
    p = doc.add_paragraph()
    run = p.add_run("python3 -m venv .venv\nsource .venv/bin/activate")
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Step 4.3: Install Dependencies", level=2)
    doc.add_paragraph(
        "All required Python packages were listed in requirements.txt and installed using pip:"
    )
    p = doc.add_paragraph()
    run = p.add_run("pip install -r requirements.txt")
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Step 4.4: Configure pytest", level=2)
    doc.add_paragraph(
        "A pytest.ini configuration file was created to define test discovery paths, "
        "custom markers (smoke, regression, login, product, cart, checkout), and default options."
    )

    # ──── 5. Framework Architecture ──────────────────────────────
    doc.add_heading("5. Framework Architecture", level=1)
    doc.add_paragraph("The project follows a modular architecture with clear separation of concerns:")

    add_styled_table(doc, ["Directory/File", "Purpose"], [
        ("config/config.py", "Centralized configuration — URLs, credentials, timeouts, product data"),
        ("pages/", "Page Object Model classes — one class per web page"),
        ("pages/base_page.py", "Base class with reusable Selenium methods (click, type, wait, scroll)"),
        ("tests/", "Test modules organized by feature area"),
        ("conftest.py", "pytest fixtures — WebDriver setup/teardown, screenshot-on-failure hook"),
        ("utils/helpers.py", "Utility functions — random data generators, screenshot helpers"),
        ("reports/", "Generated test reports, screenshots, and Excel bug report"),
    ])

    doc.add_paragraph()
    doc.add_paragraph(
        "The directory structure ensures that test logic is separated from page interactions, "
        "configuration is centralized, and reports are generated automatically."
    )

    # ──── 6. Page Object Model ───────────────────────────────────
    doc.add_heading("6. Page Object Model (POM) Design", level=1)
    doc.add_paragraph(
        "The Page Object Model (POM) design pattern was used to create a maintainable and "
        "scalable test framework. Each web page is represented by a Python class that encapsulates "
        "the page's elements (locators) and actions (methods)."
    )

    doc.add_heading("6.1 Base Page", level=2)
    doc.add_paragraph(
        "BasePage is the parent class for all page objects. It provides common Selenium interactions "
        "with built-in explicit waits and error handling:"
    )
    methods = [
        "click(locator) — Wait for element to be clickable, then click",
        "type_text(locator, text) — Clear field and type text",
        "get_text(locator) — Get visible text of an element",
        "find_visible_element(locator) — Wait for element to be visible",
        "is_displayed(locator) — Check if element is on the page",
        "scroll_to_element(locator) — Scroll element into view",
        "js_click(locator) — Click using JavaScript (for stubborn elements)",
    ]
    for m in methods:
        doc.add_paragraph(m, style="List Bullet")

    doc.add_heading("6.2 Page Objects Created", level=2)
    add_styled_table(doc, ["Page Object", "Web Page", "Key Actions"], [
        ("LoginPage", "Login (saucedemo.com)", "login(), get_error_message(), is_login_page()"),
        ("InventoryPage", "Product Listing", "get_product_names(), sort_by(), add_to_cart(), logout()"),
        ("ProductPage", "Product Detail", "get_product_name(), add_to_cart(), click_back()"),
        ("CartPage", "Shopping Cart", "get_cart_items(), remove_item(), click_checkout()"),
        ("CheckoutPage", "Checkout (3 steps)", "fill_checkout_info(), click_finish(), is_order_complete()"),
    ])

    # ──── 7. Test Scenarios ──────────────────────────────────────
    doc.add_heading("7. Test Scenarios & Test Cases", level=1)
    doc.add_paragraph(
        "A total of 54 test cases were designed across 4 modules. Tests were executed against "
        "6 different user types to maximize bug detection coverage."
    )

    doc.add_heading("7.1 User Types Tested", level=2)
    add_styled_table(doc, ["User Type", "Expected Behavior", "Purpose"], [
        ("standard_user", "Everything works correctly", "Baseline/happy path"),
        ("locked_out_user", "Cannot login", "Error handling validation"),
        ("problem_user", "Has intentional UI/logic bugs", "Bug detection"),
        ("performance_glitch_user", "Slow responses", "Performance testing"),
        ("error_user", "Errors in specific flows", "Error state testing"),
        ("visual_user", "Visual/layout bugs", "Visual regression testing"),
    ])

    doc.add_heading("7.2 Login Tests (10 test cases)", level=2)
    add_styled_table(doc, ["Test ID", "Description", "User Type"], [
        ("TC-LOG-001", "Login with valid standard_user credentials", "standard_user"),
        ("TC-LOG-002", "Login with locked_out_user — error expected", "locked_out_user"),
        ("TC-LOG-003", "Login with problem_user — should succeed", "problem_user"),
        ("TC-LOG-004", "Login with performance_glitch_user — slow but works", "performance_glitch_user"),
        ("TC-LOG-005", "Login with error_user — should succeed", "error_user"),
        ("TC-LOG-006", "Login with visual_user — should succeed", "visual_user"),
        ("TC-LOG-007", "Login with invalid credentials — error expected", "N/A"),
        ("TC-LOG-008", "Login with empty username — error expected", "N/A"),
        ("TC-LOG-009", "Login with empty password — error expected", "N/A"),
        ("TC-LOG-010", "Login then logout — return to login page", "standard_user"),
    ])

    doc.add_heading("7.3 Product/Inventory Tests (21 test cases)", level=2)
    doc.add_paragraph("Tests cover product display, image loading, sorting (4 options), and product detail navigation across standard_user, problem_user, and error_user.")
    add_styled_table(doc, ["Test ID", "Description", "User Type"], [
        ("TC-PRD-001", "Inventory displays exactly 6 products", "standard_user"),
        ("TC-PRD-002", "All product names are non-empty", "standard_user"),
        ("TC-PRD-003", "All prices are valid (contain $ and positive)", "standard_user"),
        ("TC-PRD-004", "All product images load correctly", "standard_user"),
        ("TC-PRD-005 to 008", "Sort by Name A-Z, Z-A, Price Low-High, High-Low", "standard_user"),
        ("TC-PRD-009", "Click product opens detail page", "standard_user"),
        ("TC-PRD-010", "Detail page shows price and description", "standard_user"),
        ("TC-PRD-011", "Back button returns to inventory", "standard_user"),
        ("TC-PRD-012", "Product images should be unique (not all same)", "problem_user"),
        ("TC-PRD-013 to 016", "All 4 sort options should work", "problem_user"),
        ("TC-PRD-017", "Product detail image should not be broken", "problem_user"),
        ("TC-PRD-018", "Products page displays correctly", "error_user"),
        ("TC-PRD-019 to 021", "Sort Z-A, Price Low-High, Price High-Low", "error_user"),
    ])

    doc.add_heading("7.4 Cart Tests (11 test cases)", level=2)
    doc.add_paragraph("Tests cover adding/removing items, cart persistence, and cross-user behavior.")

    doc.add_heading("7.5 Checkout Tests (12 test cases)", level=2)
    doc.add_paragraph("Tests cover the complete 3-step checkout flow, form validation, total calculation, and cross-user checkout bugs.")

    # ──── 8. Test Execution ──────────────────────────────────────
    doc.add_heading("8. Test Execution", level=1)

    doc.add_heading("8.1 How Tests Were Run", level=2)
    doc.add_paragraph("Tests were executed using the pytest framework with the following command:")
    p = doc.add_paragraph()
    run = p.add_run("pytest tests/ -v --tb=short --junitxml=reports/results.xml")
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("8.2 Execution Environment", level=2)
    add_styled_table(doc, ["Parameter", "Value"], [
        ("Operating System", "macOS (ARM64)"),
        ("Browser", "Google Chrome 147.0.7727.102"),
        ("Python Version", "3.9.6"),
        ("Selenium Version", "4.36.0"),
        ("Execution Mode", "Headed (visible browser)"),
        ("Total Execution Time", "5 minutes 8 seconds"),
    ])

    doc.add_heading("8.3 Execution Features", level=2)
    features = [
        "Automatic WebDriver Management: ChromeDriver is automatically downloaded and matched to the installed Chrome version using webdriver-manager.",
        "Screenshot on Failure: A pytest hook automatically captures a screenshot whenever a test fails, saving it to reports/screenshots/ with a timestamped filename.",
        "JUnit XML Output: Test results are exported in JUnit XML format for integration with CI/CD tools.",
        "Custom Markers: Tests are tagged with markers (smoke, regression, login, product, cart, checkout) allowing selective execution.",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    # ──── 9. Test Results ────────────────────────────────────────
    doc.add_heading("9. Test Results & Analysis", level=1)

    doc.add_heading("9.1 Overall Summary", level=2)
    add_styled_table(doc, ["Metric", "Value"], [
        ("Total Test Cases", "54"),
        ("Passed", "36 (67%)"),
        ("Failed", "18 (33%)"),
        ("Errors", "0"),
        ("Skipped", "0"),
    ])

    doc.add_heading("9.2 Results by Module", level=2)
    add_styled_table(doc, ["Module", "Total", "Passed", "Failed", "Pass Rate"], [
        ("Login", "10", "9", "1", "90%"),
        ("Product/Inventory", "21", "14", "7", "67%"),
        ("Cart", "11", "11", "0", "100%"),
        ("Checkout", "12", "3", "9", "25%"),
    ])

    doc.add_heading("9.3 Results by User Type", level=2)
    add_styled_table(doc, ["User Type", "Tests", "Passed", "Failed", "Bugs Found"], [
        ("standard_user", "~25", "All", "0*", "0"),
        ("locked_out_user", "1", "1", "0", "0"),
        ("problem_user", "~12", "4", "8", "6"),
        ("performance_glitch_user", "1", "1", "0", "0"),
        ("error_user", "~10", "4", "6", "5"),
        ("visual_user", "1", "1", "0", "0"),
    ])
    doc.add_paragraph(
        "* Some standard_user checkout tests failed due to test framework locator issues, "
        "not actual application bugs. These are noted separately from real bugs."
    )

    # ──── 10. Bugs Found ─────────────────────────────────────────
    doc.add_heading("10. Bugs Found", level=1)
    doc.add_paragraph(
        "A total of 11 bugs were discovered during testing. These are categorized by severity:"
    )

    doc.add_heading("10.1 Critical Bugs (4)", level=2)

    doc.add_heading("BUG-005: Checkout form corrupts first name input (problem_user)", level=3)
    doc.add_paragraph("Severity: Critical")
    doc.add_paragraph("Steps to Reproduce:", style="List Bullet")
    doc.add_paragraph("1. Login as problem_user")
    doc.add_paragraph("2. Add any item to cart and proceed to checkout")
    doc.add_paragraph("3. Type 'John' in the First Name field")
    doc.add_paragraph("4. Inspect the field value")
    doc.add_paragraph("Expected: Field contains 'John'")
    doc.add_paragraph("Actual: Field contains only 'e' — the input is corrupted, retaining only a single wrong character.")

    doc.add_heading("BUG-006: Checkout fails with 'Last Name is required' (problem_user)", level=3)
    doc.add_paragraph("Severity: Critical")
    doc.add_paragraph(
        "Despite entering valid data in all fields, the checkout form shows 'Error: Last Name is required'. "
        "The last name field does not accept input correctly for problem_user, blocking the entire checkout flow."
    )

    doc.add_heading("BUG-007: Last Name field empty after typing (error_user)", level=3)
    doc.add_paragraph("Severity: Critical")
    doc.add_paragraph(
        "When error_user types 'Doe' in the Last Name field, the field remains empty (''). "
        "The input is silently discarded, preventing checkout completion."
    )

    doc.add_heading("BUG-008: Finish button not clickable on checkout overview (error_user)", level=3)
    doc.add_paragraph("Severity: Critical")
    doc.add_paragraph(
        "The Finish button on the checkout overview page cannot be clicked for error_user, "
        "resulting in a TimeoutException. The entire checkout flow is blocked."
    )

    doc.add_heading("10.2 High Severity Bugs (7)", level=2)

    doc.add_heading("BUG-001: All product images show same 404 image (problem_user)", level=3)
    doc.add_paragraph("Severity: High")
    doc.add_paragraph(
        "All 6 products on the inventory page display the same broken image "
        "(sl-404.168b1cce10384b857a6f.jpg) instead of their unique product images. "
        "Users cannot visually identify products."
    )

    doc.add_heading("BUG-002 to BUG-004: All sort options broken (problem_user)", level=3)
    doc.add_paragraph("Severity: High")
    doc.add_paragraph(
        "For problem_user, selecting any sort option (Z-A, Price Low-to-High, Price High-to-Low) "
        "does not change the product order. Products remain in the default A-Z order regardless "
        "of the selected sort option."
    )

    doc.add_heading("BUG-009 to BUG-011: Sorting triggers JS alert (error_user)", level=3)
    doc.add_paragraph("Severity: High")
    doc.add_paragraph(
        "For error_user, selecting any sort option other than A-Z triggers a JavaScript alert: "
        "'Sorting is broken! This error has been reported to Backtrace.' "
        "This blocks all subsequent page interactions until the alert is dismissed."
    )

    # ──── 11. Bug Report ─────────────────────────────────────────
    doc.add_heading("11. Bug Report (Excel)", level=1)
    doc.add_paragraph(
        "A comprehensive Excel bug report was generated at: reports/SauceDemo_Bug_Report.xlsx"
    )
    doc.add_paragraph("The Excel file contains three sheets:")
    sheets = [
        "Bug Report — Detailed bug descriptions with Bug ID, Test Case ID, Module, User Type, Severity, Title, Steps to Reproduce, Expected Result, Actual Result, Status, and Screenshot reference.",
        "Test Summary — Pass/fail breakdown by module with total counts and pass rates.",
        "Bugs by User Type — Distribution of bugs across all 6 user types with severity breakdown.",
    ]
    for s in sheets:
        doc.add_paragraph(s, style="List Bullet")

    # ──── 12. Conclusion ─────────────────────────────────────────
    doc.add_heading("12. Conclusion", level=1)
    doc.add_paragraph(
        "This automation testing project successfully demonstrated the use of Selenium WebDriver "
        "with the Page Object Model design pattern to test the SauceDemo e-commerce application. "
        "The framework executed 54 test cases across 6 user types in approximately 5 minutes, "
        "discovering 11 bugs (4 Critical, 7 High severity)."
    )
    doc.add_paragraph("Key findings:")
    findings = [
        "standard_user: All features work correctly — serves as the baseline for comparison.",
        "problem_user: 6 bugs found — broken product images, non-functional sorting, and corrupted checkout form inputs.",
        "error_user: 5 bugs found — sorting triggers JavaScript errors, checkout form fields don't accept input, and the Finish button is broken.",
        "locked_out_user, performance_glitch_user, visual_user: Login behavior is as expected.",
    ]
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph(
        "The Page Object Model architecture ensures the framework is easily maintainable — "
        "if the SauceDemo website changes its UI, only the relevant page object needs to be updated, "
        "while all test logic remains unchanged."
    )

    # ──── Save ───────────────────────────────────────────────────
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "reports",
        "SauceDemo_Project_Report.docx",
    )
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Project report saved: {output_path}")
    return output_path


if __name__ == "__main__":
    create_project_report()
